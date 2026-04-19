from __future__ import annotations

import io
import re

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_runs(paragraph, text: str, default_bold: bool = False) -> None:
    parts = re.split(r"\*\*", text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.bold = default_bold or (i % 2 == 1)


def _set_margins(doc: Document, inches: float) -> None:
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


def md_to_docx(md_text: str) -> bytes:
    doc = Document()
    _set_margins(doc, 0.85)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h1_style = doc.styles["Heading 1"]
    h1_style.font.name = "Calibri"
    h1_style.font.size = Pt(13)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)

    first_h1 = True

    for raw_line in md_text.splitlines():
        line = raw_line.strip()

        if not line or re.match(r"^<!--.*-->$", line):
            continue

        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            _add_runs(p, line[4:])

        elif line.startswith("## "):
            p = doc.add_heading(line[3:], level=1)
            p.style = doc.styles["Heading 1"]

        elif line.startswith("# "):
            text = line[2:]
            if first_h1:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(18)
                run.font.name = "Calibri"
                first_h1 = False
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _add_runs(p, text, default_bold=True)

        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, line[2:])

        else:
            p = doc.add_paragraph()
            _add_runs(p, line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
